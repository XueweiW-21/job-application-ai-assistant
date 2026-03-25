"""
Generate an ATS-friendly .docx resume from a tailored markdown resume.

ATS design decisions:
  - Name and contact live in the document BODY (no Word header/footer).
  - Word header/footer is cleared entirely — name never repeats on page 2.
  - Section names use standard ATS keywords (Publications, not Select Publications).
  - Role header lines have keep_with_next=True to prevent orphaned headers.
  - No tables, text boxes, or columns — plain paragraphs only.

Usage:
    python generate_docx.py <input.md> <output.docx>

Requires:
    pip install python-docx pyyaml
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("ERROR: pyyaml is not installed. Run: pip install pyyaml")
    sys.exit(1)


# ── Load configuration from profile.yml ──────────────────────────────────────

def load_profile():
    """Load user profile from profile.yml in the project root."""
    # Walk up from script location to find project root (where profile.yml lives)
    script_dir = Path(__file__).resolve().parent
    for parent in [script_dir] + list(script_dir.parents):
        profile_path = parent / "profile.yml"
        if profile_path.exists():
            with open(profile_path, "r", encoding="utf-8") as f:
                return yaml.safe_load(f)
    print("WARNING: profile.yml not found. Using defaults.")
    return {}


PROFILE       = load_profile()
CONTACT_LINKS = PROFILE.get("links", {})

# Look for a .docx template in the templates directory
_TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
_TEMPLATE_FILES = list(_TEMPLATES_DIR.glob("*.docx")) if _TEMPLATES_DIR.exists() else []
TEMPLATE_PATH = _TEMPLATE_FILES[0] if _TEMPLATE_FILES else None


# ── Load styles from resume_styles.yml ───────────────────────────────────────

def _load_styles():
    """Load formatting styles from resume_styles.yml if it exists."""
    styles_path = _TEMPLATES_DIR / "resume_styles.yml" if _TEMPLATES_DIR.exists() else None
    if styles_path and styles_path.exists():
        with open(styles_path, "r", encoding="utf-8") as f:
            data = yaml.safe_load(f) or {}
        return data.get("styles", {}), data.get("page", {})
    return {}, {}


_STYLES, _PAGE = _load_styles()


def _s(element, prop, default=None):
    """Look up a style property, falling back to default."""
    return _STYLES.get(element, {}).get(prop, default)


# Resolved constants: styles override profile, profile overrides built-in defaults
FONT_NAME = _s("name", "font_name") or PROFILE.get("font", "Times New Roman")


# ── Core run factory ──────────────────────────────────────────────────────────

def run(paragraph, text, bold=False, italic=False, size=11.0):
    """Add a run. Always sets font name and size explicitly."""
    r = paragraph.add_run(text)
    r.font.name  = FONT_NAME
    r.font.size  = Pt(size)
    r.bold       = bold
    r.italic     = italic
    return r


def add_hyperlink(paragraph, text, url, size=10.0):
    """
    Add a hyperlink run to a paragraph.
    Renders as blue underlined text linking to url.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    rPr.append(szCs)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")  # standard Word hyperlink blue
    rPr.append(color)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    new_run.append(t)

    hl.append(new_run)
    paragraph._p.append(hl)


# ── XML helpers ───────────────────────────────────────────────────────────────

def add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(pBdr)


def add_right_tab(paragraph, pos_twips=9360):
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl"):
            body.remove(child)


def clear_headers_footers(doc):
    """Remove all header/footer content so name never repeats on page 2."""
    for section in doc.sections:
        section.different_first_page_header_footer = False
        for hdr in (section.header, section.footer,
                    section.first_page_header, section.first_page_footer,
                    section.even_page_header, section.even_page_footer):
            try:
                for p in hdr.paragraphs:
                    p.clear()
            except Exception:
                pass


# ── Inline markdown renderer ──────────────────────────────────────────────────

def _render_inline(paragraph, text, default_bold=False, size=11.0):
    """Render **bold** and *italic* markdown markers into runs."""
    parts = re.split(r"(\*\*.*?\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run(paragraph, part[2:-2], bold=True, size=size)
        elif part.startswith("*") and part.endswith("*"):
            run(paragraph, part[1:-1], italic=True, size=size)
        elif part:
            run(paragraph, part, bold=default_bold, size=size)


# ── Header block (name + contact in body) ────────────────────────────────────

def _extract_link_label(token):
    """Return display text from a markdown [text](url) token, or the token itself."""
    m = re.match(r'\[([^\]]+)\]\([^)]+\)', token.strip())
    return m.group(1) if m else token.strip()


def add_name_contact_block(doc, md_text):
    """
    Render name and contact line as body paragraphs (location line omitted).
    Known link labels from profile.yml are rendered as clickable hyperlinks.
    """
    name    = ""
    contact = ""

    for line in md_text.split("\n"):
        s = line.strip()
        if s.startswith("# ") and not s.startswith("## "):
            name = s[2:].strip()
        elif not contact and "@" in s and not s.startswith("#"):
            contact = s
            break

    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Emu(_s("name", "space_before", 0))
        p.paragraph_format.space_after  = Emu(_s("name", "space_after", 25400))
        run(p, name, bold=True, size=_s("name", "font_size", 20.0))

    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Emu(0)
        p.paragraph_format.space_after  = Emu(_s("contact", "space_after", 50800))
        contact_size = _s("contact", "font_size", 10.0)

        tokens = [t.strip() for t in contact.split("|")]
        for i, token in enumerate(tokens):
            if i > 0:
                run(p, " | ", size=contact_size)
            label = _extract_link_label(token)
            if label in CONTACT_LINKS:
                add_hyperlink(p, label, CONTACT_LINKS[label], size=contact_size)
            else:
                run(p, label, size=contact_size)


# ── Paragraph builders ────────────────────────────────────────────────────────

def add_section_header(doc, title):
    # Normalize: strip "Select"/"Selected" so header is always clean (e.g. "PUBLICATIONS")
    display = re.sub(r"^select(?:ed)?\s+", "", title, flags=re.IGNORECASE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Emu(_s("section_header", "space_before", 76200))
    p.paragraph_format.space_after  = Emu(_s("section_header", "space_after", 38100))
    run(p, display.upper(), bold=True, size=_s("section_header", "font_size", 11.0))
    if _s("section_header", "bottom_border", True):
        add_bottom_border(p)


def add_role_line(doc, company_text, date_text, is_first=False):
    p = doc.add_paragraph()
    first_sb = _s("role_header", "first_space_before", 88900)
    normal_sb = _s("role_header", "space_before", 76200)
    p.paragraph_format.space_before    = Emu(first_sb) if is_first else Emu(normal_sb)
    p.paragraph_format.space_after     = Emu(_s("role_header", "space_after", 38100))
    p.paragraph_format.keep_with_next  = True
    add_right_tab(p, pos_twips=_s("role_header", "right_tab_position", 9360))
    run(p, company_text, bold=True, size=_s("role_header", "font_size", 11.0))
    if date_text:
        run(p, "\t")
        run(p, date_text, size=_s("role_header", "font_size", 11.0))


def add_subrole_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before   = Emu(_s("job_title", "space_before", 38100))
    p.paragraph_format.space_after    = Emu(_s("job_title", "space_after", 19050))
    p.paragraph_format.keep_with_next = True
    _render_inline(p, text, size=_s("job_title", "font_size", 11.0))


def add_bullet(doc, text, is_last=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Emu(_s("bullet", "left_indent", 228600))
    p.paragraph_format.first_line_indent = Emu(_s("bullet", "first_line_indent", -114300))
    p.paragraph_format.space_before      = Emu(_s("bullet", "space_before", 0))
    last_sa = _s("bullet", "last_space_after", 38100)
    normal_sa = _s("bullet", "space_after", 0)
    p.paragraph_format.space_after       = Emu(last_sa) if is_last else Emu(normal_sa)
    bullet_char = _s("bullet", "bullet_char", "\u25cf")
    bullet_size = _s("bullet", "bullet_font_size", 6.0)
    run(p, bullet_char + "  ", size=bullet_size)
    _render_inline(p, text, size=_s("bullet", "font_size", 11.0))


def add_institution_line(doc, text, date_text=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before   = Emu(_s("institution", "space_before", 0))
    p.paragraph_format.space_after    = Emu(_s("institution", "space_after", 0))
    p.paragraph_format.keep_with_next = True
    inst_size = _s("institution", "font_size", 11.0)
    if date_text:
        add_right_tab(p, pos_twips=_s("institution", "right_tab_position", 9360))
    _render_inline(p, text, default_bold=True, size=inst_size)
    if date_text:
        run(p, "\t")
        run(p, date_text, size=inst_size)


def add_indented_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Emu(_s("education_detail", "left_indent",
                                              _s("bullet", "left_indent", 228600)))
    p.paragraph_format.space_before = Emu(0)
    p.paragraph_format.space_after  = Emu(0)
    _render_inline(p, text, size=_s("education_detail", "font_size", 11.0))


def add_skills_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Emu(_s("skills_line", "space_before", 0))
    p.paragraph_format.space_after  = Emu(_s("skills_line", "space_after", 0))
    skills_size = _s("skills_line", "font_size", 11.0)
    m = re.match(r"(\*\*.*?\*\*:?)\s*(.*)", text)
    if m:
        label = m.group(1).strip("*").rstrip(":")
        items = m.group(2)
        run(p, label + ": ", bold=True, size=skills_size)
        if items:
            run(p, items, size=skills_size)
    else:
        run(p, text, size=skills_size)


def add_publication(doc, citation, summary=None):
    pub_size = _s("publication", "font_size", 11.0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Emu(_s("publication", "space_before", 0))
    p.paragraph_format.space_after  = Emu(_s("publication", "space_after", 0))
    _render_inline(p, citation, size=pub_size)

    if summary:
        ps = doc.add_paragraph()
        ps.paragraph_format.space_before = Emu(0)
        ps.paragraph_format.space_after  = Emu(0)
        run(ps, summary, size=pub_size)

    blank = doc.add_paragraph()
    blank.paragraph_format.space_before = Emu(0)
    blank.paragraph_format.space_after  = Emu(0)


# ── Section renderers ─────────────────────────────────────────────────────────

def _strip_education_markdown(text):
    """Strip markdown syntax that should not appear in education lines."""
    # Strip ### headers
    text = re.sub(r"^#{1,4}\s+", "", text)
    # Strip wrapping *italic* markers (e.g. *Specialization: ...*)
    text = re.sub(r"^\*(.+)\*$", r"\1", text)
    # Strip wrapping **bold** markers
    text = re.sub(r"^\*\*(.+)\*\*$", r"\1", text)
    return text


def render_education(doc, lines):
    for line in lines:
        s = _strip_education_markdown(line.strip())
        if not s:
            continue
        if line.startswith("  ") or line.startswith("\t") or s.startswith("-"):
            text = s.lstrip("- ").strip()
            add_indented_line(doc, text)
        elif "|" in s:
            parts = s.split("|", 1)
            add_institution_line(doc, parts[0].strip(), date_text=parts[1].strip())
        else:
            add_institution_line(doc, s)


def render_skills(doc, lines):
    for line in lines:
        s = line.strip()
        if not s or s.startswith("---") or s.startswith("###"):
            continue
        if s.startswith("**") or "**" in s:
            add_skills_line(doc, s)
        elif s.startswith("-"):
            add_skills_line(doc, s.lstrip("- "))
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            run(p, s)


def render_experience(doc, lines):
    roles   = []
    current = None
    for line in lines:
        s = line.strip()
        if not s or s.startswith("---"):
            continue
        if s.startswith("### "):
            if current:
                roles.append(current)
            current = {"header": s[4:].strip(), "subtitle": None, "bullets": []}
        elif current is not None:
            if s.startswith("*") and s.endswith("*") and not s.startswith("**"):
                current["subtitle"] = s.strip("*").strip()
            elif s.startswith("- "):
                current["bullets"].append(s[2:].strip())
    if current:
        roles.append(current)

    for i, role in enumerate(roles):
        header    = role["header"]
        date_text = ""
        if "|" in header:
            parts       = header.rsplit("|", 1)
            header_text = parts[0].strip()
            date_text   = parts[1].strip()
        else:
            header_text = header

        add_role_line(doc, header_text, date_text, is_first=(i == 0))
        if role["subtitle"]:
            add_subrole_line(doc, role["subtitle"])
        bullets = role["bullets"]
        for j, bullet in enumerate(bullets):
            add_bullet(doc, bullet, is_last=(j == len(bullets) - 1))


def render_publications(doc, lines):
    citation = None
    summary  = None

    def flush():
        if citation:
            add_publication(doc, citation, summary)

    for line in lines:
        s = line.strip()
        if not s or s.startswith("---"):
            continue
        if re.match(r"^\d+\.\s", s):
            flush()
            citation = re.sub(r"^\d+\.\s+", "", s)
            summary  = None
        elif citation and summary is None and (
            (s.startswith("(") and s.endswith(")"))
            or (not re.match(r"^\d+\.\s", s)
                and not s.startswith("##")
                and len(s) < 200)
        ):
            # Treat the next non-empty line after a citation as its summary,
            # whether in (parens) or as a plain sentence.
            summary = s
        else:
            flush()
            citation = s
            summary  = None
    flush()


def render_section(doc, name, lines):
    add_section_header(doc, name)
    n = re.sub(r"^select(?:ed)?\s+", "", name.upper().strip())
    if n == "EDUCATION":
        render_education(doc, lines)
    elif n in ("SKILLS SUMMARY", "SKILLS"):
        render_skills(doc, lines)
    elif n == "EXPERIENCE":
        render_experience(doc, lines)
    elif "PUBLICATION" in n:
        render_publications(doc, lines)
    else:
        for line in lines:
            s = line.strip()
            if s:
                p = doc.add_paragraph()
                run(p, s)


# ── Markdown parser ───────────────────────────────────────────────────────────

def parse_resume_md(md_text):
    sections        = []
    current_section = None
    current_lines   = []

    def flush():
        if current_section is not None:
            sections.append({"section": current_section, "lines": list(current_lines)})

    for raw_line in md_text.split("\n"):
        # Strip blockquote prefix ("> " or ">") so both plain and blockquote
        # markdown formats are handled correctly by all section renderers.
        line = re.sub(r"^>\s?", "", raw_line)
        s = line.strip()
        if s.startswith("# ") and not s.startswith("## "):
            continue
        if ("Email" in s or "Tel" in s or "@" in s) and not s.startswith("##") and not s.startswith("#"):
            continue
        if not s.startswith("#") and not s.startswith("##") and current_section is None and s and not s.startswith("---"):
            continue
        if s.startswith("## "):
            flush()
            current_section = s[3:].strip()
            current_lines   = []
        else:
            current_lines.append(line)

    flush()
    return sections


# ── ATS section ordering ──────────────────────────────────────────────────────

_ATS_ORDER = ["education", "skills summary", "skills", "experience",
              "publications", "select publications"]


def sort_sections_for_ats(sections):
    def key(s):
        n = re.sub(r"^select\s+", "", s["section"].lower().strip())
        for i, label in enumerate(_ATS_ORDER):
            if n == label or n.startswith(label):
                return i
        return 99
    return sorted(sections, key=key)


# ── Entry point ───────────────────────────────────────────────────────────────

def generate_docx(md_path, output_path):
    md_text = Path(md_path).read_text(encoding="utf-8")

    if TEMPLATE_PATH and TEMPLATE_PATH.exists():
        doc = Document(str(TEMPLATE_PATH))
        clear_headers_footers(doc)
        clear_body(doc)
    else:
        doc = Document()

    # Apply page margins from styles (template already has them, but this
    # ensures consistency when no template is used or styles override)
    if _PAGE:
        sec = doc.sections[0]
        sec.top_margin    = Emu(_PAGE.get("top_margin", 914400))
        sec.bottom_margin = Emu(_PAGE.get("bottom_margin", 914400))
        sec.left_margin   = Emu(_PAGE.get("left_margin", 914400))
        sec.right_margin  = Emu(_PAGE.get("right_margin", 914400))

    add_name_contact_block(doc, md_text)

    for section in sort_sections_for_ats(parse_resume_md(md_text)):
        render_section(doc, section["section"], section["lines"])

    doc.save(output_path)
    print(f"Resume saved: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python generate_docx.py <input.md> <output.docx>")
        sys.exit(1)
    generate_docx(sys.argv[1], sys.argv[2])
