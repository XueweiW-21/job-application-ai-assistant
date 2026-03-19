"""
Tests for generate_docx.py

Run with:
    pip install -r requirements.txt pytest
    pytest tests/
"""

import sys
import re
import tempfile
from pathlib import Path

import pytest

# Add the script directory to the path so we can import the module directly
sys.path.insert(0, str(Path(__file__).parent.parent / ".claude" / "skills" / "resume-tailor" / "scripts"))
import generate_docx as gd


# ── Fixtures ──────────────────────────────────────────────────────────────────

MICHAEL_SCOTT_MD = (
    Path(__file__).parent.parent / "materials" / "examples" / "resume_master.md"
).read_text(encoding="utf-8")

MINIMAL_MD = """\
# Test Person

Email: test@example.com | Tel: (555) 000-0001
New York, NY

## Education

**Some University** | 2020
- B.S. Computer Science

## Skills Summary

**Languages:** Python (Expert), SQL (Proficient)

## Experience

### Acme Corp | Jan 2021 - Present
*Software Engineer*
- Built a thing that did stuff, improving throughput by 30%
- Worked with a team of 5 engineers across 2 time zones

## Publications

1. Test, A. (2023). "A Paper Title." *Some Journal*, 10, 12345.
(Sole author)
"""


# ── parse_resume_md ───────────────────────────────────────────────────────────

class TestParseResumeMd:
    def test_returns_list_of_sections(self):
        sections = gd.parse_resume_md(MINIMAL_MD)
        assert isinstance(sections, list)
        assert len(sections) > 0

    def test_section_keys(self):
        sections = gd.parse_resume_md(MINIMAL_MD)
        for s in sections:
            assert "section" in s
            assert "lines" in s

    def test_expected_section_names(self):
        sections = gd.parse_resume_md(MINIMAL_MD)
        names = [s["section"] for s in sections]
        assert "Education" in names
        assert "Skills Summary" in names
        assert "Experience" in names
        assert "Publications" in names

    def test_name_header_excluded(self):
        # The H1 name line should not become a section
        sections = gd.parse_resume_md(MINIMAL_MD)
        names = [s["section"] for s in sections]
        assert "Test Person" not in names

    def test_contact_line_excluded(self):
        # The contact/email line should not appear as a section
        sections = gd.parse_resume_md(MINIMAL_MD)
        names = [s["section"] for s in sections]
        assert not any("@" in n for n in names)

    def test_michael_scott_sections(self):
        sections = gd.parse_resume_md(MICHAEL_SCOTT_MD)
        names = [s["section"] for s in sections]
        assert "Education" in names
        assert "Skills Summary" in names
        assert "Experience" in names
        assert "Publications" in names

    def test_experience_lines_contain_bullets(self):
        sections = gd.parse_resume_md(MINIMAL_MD)
        exp = next(s for s in sections if s["section"] == "Experience")
        bullet_lines = [l for l in exp["lines"] if l.strip().startswith("- ")]
        assert len(bullet_lines) >= 1

    def test_empty_string_returns_empty_list(self):
        sections = gd.parse_resume_md("")
        assert sections == []


# ── sort_sections_for_ats ─────────────────────────────────────────────────────

class TestSortSectionsForAts:
    def _make_sections(self, names):
        return [{"section": n, "lines": []} for n in names]

    def test_education_comes_first(self):
        sections = self._make_sections(["Experience", "Education", "Skills Summary"])
        sorted_sections = gd.sort_sections_for_ats(sections)
        assert sorted_sections[0]["section"] == "Education"

    def test_experience_after_skills(self):
        sections = self._make_sections(["Experience", "Skills Summary", "Education"])
        sorted_sections = gd.sort_sections_for_ats(sections)
        names = [s["section"] for s in sorted_sections]
        assert names.index("Skills Summary") < names.index("Experience")

    def test_publications_after_experience(self):
        sections = self._make_sections(["Publications", "Experience", "Education"])
        sorted_sections = gd.sort_sections_for_ats(sections)
        names = [s["section"] for s in sorted_sections]
        assert names.index("Experience") < names.index("Publications")

    def test_select_publications_treated_same_as_publications(self):
        sections = self._make_sections(["Select Publications", "Experience", "Education"])
        sorted_sections = gd.sort_sections_for_ats(sections)
        names = [s["section"] for s in sorted_sections]
        assert names.index("Experience") < names.index("Select Publications")

    def test_unknown_sections_go_to_end(self):
        sections = self._make_sections(["Custom Section", "Education"])
        sorted_sections = gd.sort_sections_for_ats(sections)
        assert sorted_sections[0]["section"] == "Education"
        assert sorted_sections[-1]["section"] == "Custom Section"

    def test_full_standard_order(self):
        sections = self._make_sections(
            ["Publications", "Experience", "Skills Summary", "Education"]
        )
        sorted_sections = gd.sort_sections_for_ats(sections)
        names = [s["section"] for s in sorted_sections]
        assert names == ["Education", "Skills Summary", "Experience", "Publications"]


# ── _extract_link_label ───────────────────────────────────────────────────────

class TestExtractLinkLabel:
    def test_markdown_link_returns_label(self):
        assert gd._extract_link_label("[LinkedIn](https://linkedin.com/in/foo)") == "LinkedIn"

    def test_plain_text_returns_as_is(self):
        assert gd._extract_link_label("test@example.com") == "test@example.com"

    def test_strips_whitespace(self):
        assert gd._extract_link_label("  [GitHub](https://github.com/foo)  ") == "GitHub"

    def test_partial_markdown_returns_as_is(self):
        # Malformed link — no closing paren — should return the token as-is
        token = "[LinkedIn](https://linkedin.com"
        assert gd._extract_link_label(token) == token.strip()


# ── Integration: generate_docx ────────────────────────────────────────────────

class TestGenerateDocx:
    def test_generates_file_from_minimal_md(self, tmp_path):
        input_md = tmp_path / "resume.md"
        input_md.write_text(MINIMAL_MD, encoding="utf-8")
        output_docx = tmp_path / "output.docx"

        gd.generate_docx(str(input_md), str(output_docx))

        assert output_docx.exists()
        assert output_docx.stat().st_size > 0

    def test_generates_file_from_michael_scott(self, tmp_path):
        input_md = tmp_path / "michael.md"
        input_md.write_text(MICHAEL_SCOTT_MD, encoding="utf-8")
        output_docx = tmp_path / "michael_resume.docx"

        gd.generate_docx(str(input_md), str(output_docx))

        assert output_docx.exists()
        assert output_docx.stat().st_size > 0

    def test_output_is_valid_docx(self, tmp_path):
        from docx import Document

        input_md = tmp_path / "resume.md"
        input_md.write_text(MINIMAL_MD, encoding="utf-8")
        output_docx = tmp_path / "output.docx"

        gd.generate_docx(str(input_md), str(output_docx))

        # Should open without raising an exception
        doc = Document(str(output_docx))
        assert len(doc.paragraphs) > 0

    def test_name_appears_in_docx(self, tmp_path):
        from docx import Document

        input_md = tmp_path / "resume.md"
        input_md.write_text(MINIMAL_MD, encoding="utf-8")
        output_docx = tmp_path / "output.docx"

        gd.generate_docx(str(input_md), str(output_docx))

        doc = Document(str(output_docx))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Test Person" in all_text

    def test_section_headers_appear_in_docx(self, tmp_path):
        from docx import Document

        input_md = tmp_path / "resume.md"
        input_md.write_text(MINIMAL_MD, encoding="utf-8")
        output_docx = tmp_path / "output.docx"

        gd.generate_docx(str(input_md), str(output_docx))

        doc = Document(str(output_docx))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "EDUCATION" in all_text
        assert "EXPERIENCE" in all_text
        assert "SKILLS SUMMARY" in all_text
